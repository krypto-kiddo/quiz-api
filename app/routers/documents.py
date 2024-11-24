from fastapi import APIRouter, HTTPException, Request, Depends, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.sql import text
from app.models import Document
from app.database import get_db
import aiofiles
import os
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument
from sqlalchemy.dialects.postgresql import TSVECTOR
from fastapi.responses import JSONResponse
from datetime import datetime
router = APIRouter()

### UPLOAD NEW FILE ###
@router.post("/api/files/upload", summary="Upload a document")
async def upload_file(request: Request, db: AsyncSession = Depends(get_db)):
    # Parse form data to extract the file
    form = await request.form()
    if len(form) != 1:  # Ensure only one file is uploaded
        raise HTTPException(status_code=400, detail="Exactly one file must be uploaded.")
    
    # Get the file (keyless)
    file = list(form.values())[0]
    if not hasattr(file, "filename"):
        raise HTTPException(status_code=400, detail="No valid file uploaded.")

    # Validate file type
    allowed_types = ["text/plain", "application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    # Save file temporarily
    temp_file_path = f"temp_{file.filename}"
    async with aiofiles.open(temp_file_path, 'wb') as out_file:
        content = await file.read()
        await out_file.write(content)

    # Extract content based on file type
    extracted_content = None
    if file.content_type == "text/plain":
        # Handle plain text
        async with aiofiles.open(temp_file_path, 'r', encoding="utf-8") as f:
            extracted_content = await f.read()
    elif file.content_type == "application/pdf":
        # Extract text from PDF
        extracted_content = extract_text(temp_file_path)
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Extract text from DOCX
        doc = DocxDocument(temp_file_path)
        extracted_content = "\n".join([p.text for p in doc.paragraphs])

    # Generate custom file ID
    query = await db.execute(select(Document.file_id))
    existing_ids = [int(file_id.replace("file", "")) for file_id, in query.fetchall()]
    next_id = max(existing_ids, default=0) + 1
    custom_file_id = f"file{str(next_id).zfill(3)}"  # e.g., file001, file002

    # Compute the search_vector (PostgreSQL function `to_tsvector`)
    search_vector = await db.scalar(
        text("SELECT to_tsvector(:text_content)")
        .bindparams(text_content=extracted_content)
    )

    # Save file details to the database
    new_document = Document(
        file_id=custom_file_id,
        name=file.filename,
        file_type=file.content_type,
        content=extracted_content,  # Store extracted text content
        search_vector=search_vector
    )
    db.add(new_document)
    await db.commit()

    # Remove the temporary file
    os.remove(temp_file_path)

    return {"message": "File uploaded successfully", "file_id": new_document.file_id}

### GET FILE FROM ID ###
@router.get("/api/documents/{file_id}", summary="Get document details by file ID")
async def get_document(file_id: str, db: AsyncSession = Depends(get_db)):
    # Query the database for the document with the given file_id
    query = await db.execute(select(Document).where(Document.file_id == file_id))
    document = query.scalars().first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Return the document details in a standardized format
    response = {
        "success": True,
        "message": "Document fetched successfully.",
        "data": {
            "file_id": document.file_id,
            "name": document.name,
            "file_type": document.file_type,
            "content": document.content,
        },
        "metadata": {
            "length": len(document.content) if document.content else 0,  # Length of content
            "retrieved_at": datetime.utcnow().isoformat(),  # Timestamp of retrieval
        },
    }
    return JSONResponse(content=response, media_type="application/json")



### SEARCH THROUGH FILES ###
@router.get("/api/files/search", summary="Search files")
async def search_files(
    q: str = Query(..., description="Search query"),
    page: int = Query(1, description="Page number"),
    limit: int = Query(10, description="Number of results per page"),
    sort: str = Query("relevance", description="Sorting order: relevance or name"),
    type: str = Query(None, description="File types to filter by, e.g., txt,pdf"),
    db: AsyncSession = Depends(get_db)
):
    # Validate sort parameter
    if sort not in ["relevance", "name"]:
        raise HTTPException(status_code=400, detail="Invalid sort parameter")

    # Prepare the offset and query parameters
    offset = (page - 1) * limit

    # Map file types to MIME types
    mime_map = {
        "txt": "text/plain",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    file_types = [mime_map[file_type.strip()] for file_type in type.split(",")] if type else []

    # Use `to_tsquery` for search functionality
    query_str = f"""
    SELECT
        file_id,
        name,
        file_type,
        ts_rank(search_vector, to_tsquery(:query)) AS relevance
    FROM documents
    WHERE search_vector @@ to_tsquery(:query)
    {f"AND file_type = ANY(:file_types)" if file_types else ""}
    ORDER BY {"relevance DESC" if sort == "relevance" else "name ASC"}
    LIMIT :limit OFFSET :offset
    """
    params = {
        "query": q,  # Pass the search query
        "limit": limit,
        "offset": offset,
        "file_types": file_types
    }

    # Execute the query
    result = await db.execute(text(query_str), params)
    rows = result.fetchall()

    # Fetch total count for pagination
    count_query_str = f"""
    SELECT COUNT(*) FROM documents
    WHERE search_vector @@ to_tsquery(:query)
    {f"AND file_type = ANY(:file_types)" if file_types else ""}
    """
    count_result = await db.execute(text(count_query_str), {"query": q, "file_types": file_types})
    total_results = count_result.scalar()

    # Calculate total pages
    total_pages = (total_results // limit) + (1 if total_results % limit > 0 else 0)

    # Format the response
    return {
        "success": True,
        "query": {
            "search_term": q,
            "file_types": file_types,
            "sort": sort,
            "page": page,
            "limit": limit
        },
        "pagination": {
            "current_page": page,
            "total_results": total_results,
            "total_pages": total_pages,
            "limit": limit
        },
        "results": [
            {
                "file_id": row.file_id,
                "name": row.name,
                "file_type": row.file_type,
                "relevance": row.relevance,
            }
            for row in rows
        ]
    }

